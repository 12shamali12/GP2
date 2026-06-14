"""
DentyHub - Individual Reflection Reports generator (academic style).

Produces five .docx files, one per team member, in the formal academic
register expected by a college reflection report. Uses:

    - Measured first-person where genuine reflection is warranted; third-
      person for objective claims and evidence.
    - Explicit citations to sections, figures, features (F#), requirements
      (FR-#/NFR-#) and test cases (TC-#) of the group report.
    - Standard technical vocabulary (ISO/IEC 25010 categories, WCAG AA,
      OWASP guidance, NestJS / Prisma / React).
    - Structured paragraphs with topic sentence, evidence, and rationale.
    - No contractions, colloquialisms, or rhetorical flourishes.

Run:
    python individual_report_generator.py
"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------


def new_doc():
    doc = Document()
    for name in ("Normal", "Heading 2", "Heading 3"):
        st = doc.styles[name]
        st.font.name = "Cambria"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Heading 2"].font.size = Pt(16)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(13)
    doc.styles["Heading 3"].font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)
    return doc


def H2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def H3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def P(doc, text, bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align


def B(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def TABLE(doc, headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0]
    for i, h in enumerate(headers):
        cell = head.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E7EEF2")
        tc_pr.append(shd)
    for r_idx, row in enumerate(rows, start=1):
        row_cells = t.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cell = row_cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_cm is not None and len(col_widths_cm) == len(headers):
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph("")


def PAGEBREAK(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def TITLE_PAGE(doc, student_name, student_id):
    for line in [
        "Jordan University of Science and Technology",
        "College of Computer Sciences and Information Technology",
        "Department of Software Engineering",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(13)
    for _ in range(3):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("DentyHub")
    trun.bold = True
    trun.font.size = Pt(26)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(
        "A Supervised Dental Training Platform Connecting Patients, "
        "Student Doctors, Clinical Supervisors, and Administrators"
    )
    srun.italic = True
    srun.font.size = Pt(12)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = p.add_run("Graduation Project Individual Reflection Report")
    pr.bold = True
    pr.font.size = Pt(15)
    for _ in range(2):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared by").italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{student_name} (Student ID: {student_id})")
    r.font.size = Pt(13)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Supervised by").italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Dr. Zakarea M. Al Shara'a").font.size = Pt(12)
    for _ in range(2):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("May 2026").bold = True
    PAGEBREAK(doc)


def UNDERTAKING(doc, student_name):
    H2(doc, "Undertaking")
    P(doc,
        "I hereby declare that the project titled \"DentyHub: A Supervised "
        "Dental Training Platform\" is an original work conducted by the "
        "undersigned team in partial fulfillment of the requirements for "
        "the Bachelor of Science degree in Software Engineering at the "
        "Department of Software Engineering, Jordan University of Science "
        "and Technology, Irbid, Jordan.")
    P(doc,
        "The lessons, observations, and individual contributions described "
        "in this reflection report represent an honest assessment of the "
        "role the undersigned personally fulfilled within the project "
        "team. The literature and external software components referenced "
        "in this work are cited in the References section of the group "
        "report. This individual report has not been submitted, in whole "
        "or in part, to any other university or academic institution.")
    doc.add_paragraph("")
    doc.add_paragraph("")
    P(doc, "Student Name: " + student_name)
    doc.add_paragraph("")
    P(doc, "Signature: ____________________________")
    doc.add_paragraph("")
    P(doc, "Date: ____ / ____ / 2026")
    PAGEBREAK(doc)


def CONTENTS(doc):
    H2(doc, "Contents")
    P(doc,
        "The reader may regenerate the page-numbered Table of Contents by "
        "placing the cursor below this line and selecting References → "
        "Table of Contents → Update Field in Microsoft Word. The TOC is "
        "produced from the Heading 2 and Heading 3 styles applied "
        "throughout this document.",
        italic=True)
    P(doc,
        "[ The Table of Contents will be rendered here upon Update Field. ]",
        italic=True)
    PAGEBREAK(doc)


# ---------------------------------------------------------------------------
# Omar's report — comprehensive, academic register
# ---------------------------------------------------------------------------


def omar_report():
    doc = new_doc()
    TITLE_PAGE(doc, "Omar Mohammed Subhi Al Shamali", "154072")
    UNDERTAKING(doc, "Omar Mohammed Subhi Al Shamali")
    CONTENTS(doc)

    # ==== Chapter 1 ====================================================
    H2(doc, "CHAPTER 1: Lessons Learned")
    P(doc,
        "This chapter documents the lessons derived from serving as "
        "project lead on DentyHub over the two academic cycles GP1 and "
        "GP2. The reflection is organised into three categories: "
        "strategies that proved consistently effective in advancing the "
        "project; areas in which the author's approach could be improved; "
        "and broader observations gained through the experience of "
        "leading a five-person software engineering team toward a "
        "deployable system.")

    H3(doc, "1.1 Strategies and Processes that led to Success")
    P(doc,
        "The strategies that contributed most to the project's progress "
        "were structural rather than feature-level. Three are discussed "
        "below, each illustrated by concrete artefacts in the codebase.")
    P(doc,
        "First, the project prioritised data-model design before "
        "user-interface design. Approximately two weeks of GP1 were "
        "devoted to the Prisma schema before any user-facing screen was "
        "constructed. This investment generated returns throughout GP2: "
        "the schema, comprising forty-one models and eighteen enumerations "
        "(Appendix D of the group report), absorbed every subsequent "
        "feature without requiring restructuring. The Appointment, "
        "CaseReport, and DoctorClinicCaseProgress relations - the central "
        "data triangle that encodes graduation credit - were specified "
        "during GP1 and remained unchanged through GP2 despite the "
        "addition of multi-supervisor reporting, the six-game arcade "
        "engine, and the cohort-progression workflow.")
    P(doc,
        "Second, the team adopted explicit written process rules in "
        "preference to informal agreements. A single-author rule for "
        "Prisma migrations was introduced following a parallel-migration "
        "merge conflict during GP1 and was documented in the project "
        "README. A translation-key audit checklist was added to the "
        "pull-request template following an Arabic localisation oversight "
        "and has subsequently identified seven additional missing keys. "
        "These rules are summarised in Appendix K (Lessons Learned) of "
        "the group report. The pattern observed is that small, written "
        "process rules are more reliable than larger unwritten "
        "agreements among team members operating under deadline pressure.")
    P(doc,
        "Third, the engagement layer was constructed using an "
        "architectural pattern that enforced modular boundaries between "
        "games. Each arcade game is contained within a full-viewport "
        "focus stage and communicates with the surrounding application "
        "only through a portal-mounted HUD slot and a small props "
        "contract (level, onFinish, onCancel, hudSlot). The six implemented "
        "games (Plaque Blaster, Tooth Defender 3D, Floss Rush, Tooth IQ, "
        "Match Lab, and Brush Buddy) share no internal dependencies, and "
        "the addition of a seventh game would require modifications to a "
        "single file. This modular boundary, while not visible to the "
        "end user, materially reduces future maintenance cost.")
    P(doc,
        "A fourth strategy worth recording is the enforcement of "
        "engagement-level rules at the data layer rather than at the "
        "application layer. The once-per-day attempt rule on the arcade "
        "is enforced via a UNIQUE constraint on (patientId, gameType, "
        "dateKey) in the ArcadeAttempt model; the daily smile check-in "
        "is enforced by an analogous constraint on SmileCheckin. A "
        "tampered client cannot bypass either rule because the violation "
        "is rejected by PostgreSQL at insertion. This pattern of "
        "encoding engagement rules in the schema rather than in the "
        "controller layer is a contribution the author intends to apply "
        "to future projects.")

    H3(doc, "1.2 Areas of potential improvement")
    P(doc,
        "Three areas of improvement merit explicit acknowledgement.")
    P(doc,
        "The first concerns scope management. The patient engagement "
        "layer - the streak feature, the badge system, and the six "
        "arcade games - was not present in the GP1 project plan. These "
        "features emerged from a brainstorming session during GP2 "
        "(documented in Section 4.1 of the group report under milestone "
        "M8) and ultimately consumed approximately four weeks of "
        "calendar time, more than originally budgeted for any single "
        "GP2 deliverable. In a future project, the author would attempt "
        "to identify engagement features earlier in the planning phase "
        "and interleave them with the core workflow rather than "
        "introduce them mid-cycle.")
    P(doc,
        "The second area concerns automated testing. The project relied "
        "on a manual scenario-test pass before each release, documented "
        "in Section 5.2 of the group report. While this approach "
        "identified the majority of defects, two regressions during "
        "week 19 of GP2 were discovered only through manual exploration "
        "and would have been detected by an automated end-to-end test "
        "covering the booking-report-decision pathway. The author "
        "underestimated the maintenance cost of an automated end-to-end "
        "suite at this scale; future projects would benefit from the "
        "introduction of a single end-to-end test during week one and "
        "gradual expansion thereafter.")
    P(doc,
        "The third area concerns task estimation. Tasks involving "
        "technologies the author had not previously deployed - shader-"
        "aware React via three.js for Tooth Defender 3D, the right-to-"
        "left internationalisation sweep, and a PostgreSQL JSON-"
        "containment query for the multi-supervisor flow - consistently "
        "required two to three times the originally estimated effort. "
        "The corrective heuristic the author intends to apply in future "
        "is to multiply initial estimates by a factor of three when the "
        "task involves a library or pattern not previously deployed. "
        "Repeated underestimation of unfamiliar work produced "
        "unnecessary schedule pressure during the final weeks of GP2.")

    H3(doc, "1.3 Other lessons learned")
    P(doc,
        "Two broader lessons should be recorded.")
    P(doc,
        "The first concerns scope discipline. Several feature ideas "
        "considered during the project - patient payment integration, "
        "SMS reminder delivery, in-clinic queue display - were "
        "technically interesting but did not advance the project's "
        "stated value proposition (Section 1.3 of the group report). "
        "Excluding such features from scope proved more difficult than "
        "implementing them would have been. The explicit out-of-scope "
        "list in Section 1.3 of the group report represents the "
        "outcome of this exercise. The lesson learned is that the "
        "decision to exclude a feature requires the same justification "
        "as the decision to include one.")
    P(doc,
        "The second concerns operational detail. Several engineering "
        "decisions - the use of an Asia/Amman date-key for daily "
        "uniqueness rather than UTC, the application of Prisma "
        "transactions to multi-write operations, the placement of "
        "AdminGuard at the controller layer rather than at the route "
        "table - have no visible effect when implemented correctly. "
        "However, the same decisions, implemented incorrectly, would "
        "produce visible product failures: streaks resetting "
        "incorrectly across midnight, partial writes leaving the "
        "database in inconsistent states, and elevated endpoints "
        "leaking data through newly added routes. The lesson is that "
        "the quality of an operational decision is best measured not "
        "by its visibility but by the visibility of its absence.")
    P(doc,
        "Finally, the experience of leading a multi-person team "
        "produced a non-technical lesson regarding the structuring of "
        "work. The most effective approach proved to be the "
        "construction of vertical slices - well-scoped, individually "
        "ownable units of work with clear boundaries - rather than "
        "horizontal layers requiring coordination across the team. The "
        "author acknowledges that this principle was not applied as "
        "consistently as warranted; in retrospect, several tasks "
        "absorbed during GP2 could have been delegated had the slicing "
        "been performed earlier.")
    PAGEBREAK(doc)

    # ==== Chapter 2 ====================================================
    H2(doc, "CHAPTER 2: Individual Contributions")
    P(doc,
        "This chapter quantifies and characterises the author's "
        "contribution to the DentyHub project. Quantitative figures "
        "are verifiable against the project repository at "
        "github.com/12shamali12/GP1 via git log analysis. Qualitative "
        "claims are cross-referenced to the corresponding sections, "
        "figures, features, requirements, and test cases of the group "
        "report.")

    H3(doc, "2.1 Amount of Work")
    P(doc,
        "As project lead, the author authored the largest single "
        "share of the codebase across both academic cycles. The "
        "following table summarises authored output by functional "
        "surface.")
    TABLE(doc,
        ["Functional Surface", "Approximate Authored LOC", "Primary Files"],
        [
            ["Project leadership and repository governance", "—",
             "Code review on every pull request; README; branch policy"],
            ["Prisma schema architecture", "approx. 1,400",
             "schema.prisma (41 models, 18 enums); 12 append-only migrations"],
            ["Backend service architecture (appointments, supervisor decisions, multi-supervisor reporting, chat, notifications, smile streak, arcade)",
             "approx. 5,200", "approx. 50 controllers, services, DTOs, and guards"],
            ["Arcade engine and six game components", "approx. 3,400",
             "6 game files; hub; intro card; focus stage; threshold logic"],
            ["Healthy Smile Streak (backend and frontend)", "approx. 1,400",
             "smile-streak module; surface; badges; widget; streak leaderboard"],
            ["Internationalisation dictionary (EN and AR with full RTL)", "approx. 3,000 keys",
             "dictionary.ts; language-provider; RTL coverage across patient screens"],
            ["Patient booking flow and appointment lifecycle interface", "approx. 1,800",
             "patient-slot-modal; booking-actions; status pills; history view"],
            ["UX/UI design system (dark theme, responsive design, reduced motion)", "approx. 2,000",
             "globals.css; design tokens; role shells; settings panel"],
            ["Administrator planning and supervisor decision interfaces", "approx. 2,400",
             "Four-tab planning workspace; supervisor queue and decision panel"],
            ["Group report, individual reports, UML and architecture diagrams",
             "8 .docx files; 7 figures",
             "report_generator.py; individual_report_generator.py; Mermaid and PlantUML sources"],
        ],
        col_widths_cm=[5.8, 4.4, 5.6],
    )
    P(doc,
        "Quantitative metrics alone do not fully characterise the "
        "contribution. The commits also encompass code review on every "
        "team-member pull request, architectural decisions that "
        "established the boundaries within which other team members "
        "operated, and substantial unscheduled effort responding to "
        "environment-setup issues encountered by team members. The "
        "figures reported in the table constitute a lower bound on "
        "total effort rather than an upper bound.")

    H3(doc, "2.2 Quality of Work")
    P(doc,
        "The author defines code quality on this project according to "
        "three measurable properties: type-correctness, testability, "
        "and explanatory clarity. Each is discussed below with "
        "reference to the corresponding evidence in the group report.")
    P(doc,
        "Type-correctness. Every pull request authored passed "
        "TypeScript strict-mode compilation and ESLint validation with "
        "zero errors prior to merge. The continuous-integration "
        "pipeline (NFR-36 and NFR-37 of the group report) enforces "
        "these requirements at the repository level. As project lead, "
        "the author applied a consistent policy that pull-request "
        "authors must resolve their own lint and compilation errors "
        "rather than merge with workarounds; this policy contributed "
        "measurably to the overall consistency of the codebase.")
    P(doc,
        "Testability. The non-obvious algorithmic components of the "
        "system - the Asia/Amman date-key arithmetic, the arcade unlock "
        "computation (computeUnlockedLevel in backend/src/arcade/"
        "arcade.service.ts), and the booking conflict-detection logic - "
        "are accompanied by unit tests. The thirty scenario tests "
        "documented in Section 5.2 of the group report constitute the "
        "end-to-end validation suite executed prior to each release. "
        "Performance metrics reported in Section 5.3 (login p95 below "
        "250 ms; dashboard first contentful paint below 1500 ms on a "
        "4G connection) confirm that NFR-2 and NFR-7 are satisfied "
        "against the seeded build.")
    P(doc,
        "Explanatory clarity. Each backend module file contains a "
        "header docstring describing its responsibility. Non-obvious "
        "schema decisions (the denormalised DoctorClinicCaseProgress "
        "table, the AppointmentEvent audit-trail model, the UNIQUE "
        "constraints encoding daily-attempt rules) are documented with "
        "inline comments specifying the design rationale. The UML "
        "diagrams in Section 3.2 of the group report were generated "
        "from the implemented schema rather than from aspirational "
        "designs. The intended outcome is that an external reviewer "
        "may open any file in the repository and reconstruct the "
        "design rationale without requiring author consultation.")

    H3(doc, "2.3 Specialized Contributions")

    P(doc, "Project leadership and architectural direction.", bold=True)
    P(doc,
        "The author established the technical direction of the project "
        "throughout both academic cycles. Specific decisions for which "
        "the author held primary responsibility include: the pnpm "
        "workspace monorepo layout; the module-per-domain convention "
        "on the backend; the feature-folder convention on the "
        "frontend; the controller-layer placement of authorisation "
        "guards (NFR-6) to prevent silent data leakage through newly "
        "added routes; the enforcement of engagement rules via "
        "database constraints rather than client-side checks; and "
        "the append-only migration policy with a single-author rule "
        "(Appendix G of the group report). The cumulative effect of "
        "these decisions is visible in the fact that every non-"
        "functional requirement in Section 2.2 is enforced either "
        "through continuous integration or through code review.")

    P(doc, "Data model design.", bold=True)
    P(doc,
        "The author designed the system schema comprising forty-one "
        "models and eighteen enumerations. The central design - the "
        "Appointment, CaseReport, and DoctorClinicCaseProgress "
        "relation triangle - ensures that graduation credit is "
        "awarded atomically: the supervisor approval transition "
        "updates both the report status and the doctor's progress "
        "record within a single Prisma transaction. This atomicity "
        "guarantee, encoded in the schema and enforced by NFR-11, "
        "constitutes the architectural property that distinguishes "
        "DentyHub from a generic appointment management system. The "
        "schema additionally encodes the daily-attempt uniqueness "
        "constraints, the partner-pair and group-supervisor "
        "relationships used by the rotation planning subsystem, and "
        "the AppointmentEvent audit trail (feature F32).")

    P(doc, "Multi-supervisor reporting workflow.", bold=True)
    P(doc,
        "The author designed and implemented the multi-supervisor "
        "report submission workflow that permits a student doctor to "
        "address a single case report to one or more clinical "
        "supervisors. The primary supervisor is recorded in the "
        "reviewerSupervisorId column of CaseReport; additional "
        "supervisors are persisted as a JSON array in the formData "
        "column and surfaced into each supervisor's queue via a "
        "PostgreSQL JSON-containment query (the array_contains "
        "operator). Any of the addressed supervisors is authorised "
        "to review the report. Notification fan-out on submission "
        "and on decision ensures that all addressed supervisors "
        "remain synchronised regarding the report state. The "
        "decision to store the additional reviewer set as a JSON "
        "array rather than introduce a relational join table was "
        "made on the basis of three considerations: implementation "
        "speed during the freeze week, the absence of any required "
        "schema migration, and the small expected cohort size per "
        "report.")

    P(doc, "Arcade engine and game implementations.", bold=True)
    P(doc,
        "The author designed the arcade engine architecture and "
        "implemented all six included games. The engine architecture "
        "comprises the focus-mode stage, the per-game props contract "
        "({ level, onFinish, onCancel, hudSlot }), the per-level "
        "threshold curves, the sticky sequential unlock algorithm, "
        "and the server-side enforcement of the once-per-day attempt "
        "rule. The implemented games include: Plaque Blaster (a "
        "thirty-second target-tap grid with an endless Level 11 "
        "mode); Tooth Defender 3D (a three.js scene comprising a "
        "tooth model, a camera, and a projectile loop, representing "
        "the author's first shader-aware React implementation); "
        "Floss Rush (a three-lane endless runner); Tooth IQ (a "
        "deck-based multiple-choice quiz with one hundred unique "
        "questions distributed across ten levels); Match Lab (a "
        "memory matching game with per-level preview windows); and "
        "Brush Buddy (a Simon-style brushing pattern recognition "
        "game). The shared engine contract permits the addition of a "
        "seventh game through the modification of a single file.")

    P(doc, "Internationalisation and right-to-left support.", bold=True)
    P(doc,
        "The author implemented the project's internationalisation "
        "infrastructure as a custom language provider rather than as "
        "an integration with an external library. The decision was "
        "made on the basis that the system's requirements - string "
        "lookup with parameter interpolation - are sufficiently "
        "narrow to be served by a custom implementation that exposes "
        "only the useTranslation hook. The translation dictionary "
        "grew from approximately seven hundred keys at the start of "
        "GP2 to approximately one thousand one hundred keys at "
        "submission, distributed across English and Arabic. The "
        "right-to-left presentation pass covered every patient-"
        "facing screen, satisfying NFR-18 (no LTR leakage). The "
        "pre-pull-request translation-key audit checklist, "
        "introduced following an Arabic localisation oversight, has "
        "subsequently identified seven additional missing keys.")

    P(doc, "Group report and UML documentation.", bold=True)
    P(doc,
        "The author authored the group project report through a "
        "Python generator (report_generator.py) that produces "
        "DentyHub_GP2_Group_Report.docx containing five chapters, "
        "the requirements table mapping seventy-eight functional "
        "and thirty-four non-functional requirements to the thirty-"
        "two features in Section 2.1, the References section, and "
        "twelve appendices. The author additionally produced the "
        "majority of the UML figures in Chapter 3 of the group "
        "report: the class diagram (Figure 3.3) generated from the "
        "Prisma schema; the sequence diagrams for patient booking "
        "and supervisor decision (Figures 3.4 and 3.5); the "
        "appointment state machine (Figure 3.6); and the streak "
        "check-in activity diagram (Figure 3.7). The author also "
        "implemented the individual-report generator that produced "
        "the present report and its four companion reports.")

    P(doc, "Patient booking flow and atomic conflict detection.", bold=True)
    P(doc,
        "The author implemented the patient booking workflow end-"
        "to-end. The frontend components include the slot browser "
        "with date, clinic, and case-category filters; the patient-"
        "slot-modal for slot selection; the booking confirmation "
        "dialog; the appointment history panel with the AWAITING_REPORT "
        "status pill (feature F9); and the colour-coded status "
        "indicators on patient appointment cards. The backend "
        "conflict-detection logic executes within a Prisma "
        "transaction: a SELECT FOR UPDATE on the slot row, followed "
        "by INSERT operations for the Appointment, the corresponding "
        "AppointmentEvent, and the supervisor notification. The "
        "row-level lock ensures that concurrent booking submissions "
        "are serialised, satisfying NFR-10 by construction. TC-08 of "
        "Section 5.2 of the group report verifies this end-to-end.")

    P(doc, "User experience and accessibility infrastructure.", bold=True)
    P(doc,
        "The author designed the project's visual identity, "
        "comprising the denty-prefixed CSS class system, the design "
        "tokens, the light and dark theme implementation via the "
        "html.dark class strategy in globals.css, and the role-"
        "shell layouts shared across all four user roles. The "
        "mobile-first responsive design ensures that every patient-"
        "facing screen renders at a 360-pixel viewport without "
        "horizontal scrolling (NFR-32). The author additionally "
        "implemented the reduced-motion preference toggle in the "
        "settings panel (use-settings-prefs.ts), which permits "
        "users with vestibular sensitivities to disable arcade "
        "animations and the streak widget's flame animation, in "
        "compliance with WCAG 2.1 Success Criterion 2.3.3 "
        "(Animation from Interactions). The @media (prefers-"
        "reduced-motion: reduce) declarations in globals.css "
        "additionally honour operating-system-level motion "
        "preferences.")

    P(doc, "Healthy Smile Streak engine.", bold=True)
    P(doc,
        "The author implemented the Healthy Smile Streak feature "
        "end-to-end. The backend module (backend/src/smile-streak) "
        "comprises the SmileCheckin model with its UNIQUE "
        "constraint on (patientId, dateKey-in-Asia/Amman); the "
        "streak computation that compares the prior day's date-key "
        "to the current day's to determine consecutive-day "
        "increments; and the automatic badge-award logic at "
        "milestones of 3, 7, 30, and 100 consecutive days. The "
        "frontend components include the daily ritual interface "
        "presenting brushing, flossing, and mouthwash tasks within "
        "a thirty-second window; the badge wall; the streak "
        "summary widget on the patient overview; and the streak "
        "leaderboard with a three-metric switcher (current streak, "
        "best-ever streak, and cumulative points). The choice of "
        "Asia/Amman date-key over UTC date-key ensures that "
        "midnight boundaries align with the user's lived "
        "experience.")

    P(doc, "Communications subsystem.", bold=True)
    P(doc,
        "The author implemented the notifications subsystem, "
        "comprising the persistence of a Notification row for each "
        "account vetting decision, supervisor decision, redo "
        "request, partner request, group join request, and chat "
        "message. The associated bell-menu interface surfaces these "
        "notifications to the recipient. The author additionally "
        "implemented the one-to-one chat workflow between role-"
        "paired users (patient and doctor, doctor and supervisor, "
        "administrator and any role), comprising participant search, "
        "unread message counting, mark-conversation-read on "
        "interaction, and optimistic-update message dispatch. The "
        "notification fan-out for the multi-supervisor reporting "
        "workflow reuses this same infrastructure.")

    P(doc, "Administrator and supervisor interfaces.", bold=True)
    P(doc,
        "The author implemented the four-tab administrator planning "
        "workspace (Resources, Plans, Assignments, and Supervisors) "
        "that consumes the CRUD endpoints contributed by the team "
        "member responsible for backend administrative endpoints. "
        "The administrator user directory with its search, filter, "
        "block, freeze, and re-approval actions, the Add-case deep-"
        "link from the Cases page into the planning Resources tab, "
        "and the administrator Settings tab were similarly "
        "implemented by the author. On the supervisor side, the "
        "author implemented the report review queue and the three-"
        "way decision panel (Approve, Needs Edit, Redo), each "
        "accepting an optional mark, rating, and feedback "
        "annotation.")

    P(doc, "Doctor workspace and rotation plan presentation.", bold=True)
    P(doc,
        "The author implemented the doctor-facing rotation plan "
        "workspace, comprising the hero card displaying the next "
        "upcoming clinic day, the schedule strip presenting the "
        "assigned clinic, supervisor, and shift template, the per-"
        "case progress visualisation indicating the categories the "
        "student must still complete to meet graduation "
        "requirements, and the daily desk view aggregating the "
        "next appointment, pending case reports, and supervisor "
        "decisions awaiting response. This workspace constitutes "
        "the primary surface accessed by student doctors upon "
        "login.")

    P(doc, "Session lifecycle and authorisation policy.", bold=True)
    P(doc,
        "While the implementation of the authentication module is "
        "the responsibility of another team member, the author "
        "established the higher-level session lifecycle policy. "
        "The chosen design - a seven-day JSON Web Token signed via "
        "passport-jwt, without refresh-token rotation - represents "
        "a deliberate engineering trade-off. Refresh-token rotation "
        "would add a server-side token persistence layer, a "
        "refresh endpoint, and a rotation strategy to defend "
        "against token replay. These additions are appropriate at "
        "production scale but introduce complexity disproportionate "
        "to the threat model of a graduation-project pilot. The "
        "seven-day token lifetime maintains a narrow authentication "
        "surface in the present implementation while preserving the "
        "architectural option to layer refresh-token rotation onto "
        "the existing endpoints when usage scale warrants it.")

    H3(doc, "2.4 Initiative and Reliability")
    P(doc,
        "Initiative on this project is documented by the set of "
        "features whose proposal and implementation were the author's "
        "responsibility. These include the arcade engine and its six "
        "games; the schema redesign that introduced the per-doctor "
        "case-progress denormalised table following supervisor "
        "interviews; the streak leaderboard two-tab switcher; the "
        "multi-supervisor reporting flow with JSON-array persistence "
        "of additional supervisors; the reduced-motion preference "
        "toggle in the settings panel; and the patient booking "
        "modal's atomic conflict-detection logic. Each of these "
        "features was proposed and implemented by the author and "
        "represents architectural or product direction beyond "
        "explicitly assigned tasks.")
    P(doc,
        "Reliability is documented by the project milestones listed "
        "in Section 1.4 of the group report. All ten milestones "
        "closed within the GP2 academic window. Milestone M8 (the "
        "initial three arcade games) extended by approximately one "
        "week due to the unfamiliar shader-aware React implementation "
        "of Tooth Defender 3D; the schedule slip was absorbed by "
        "reducing the Match Lab difficulty curve from four tiers to "
        "three in milestone M9, a reduction subsequently validated "
        "by play-testing.")
    P(doc,
        "Additional indicators of reliability not present in the "
        "commit history include consistent attendance at the weekly "
        "supervisor check-ins with concrete progress demonstrations; "
        "same-day response to team-member queries via the project's "
        "communication channels; and the consistent provision of "
        "explanatory feedback rather than mere rejection during "
        "code review.")
    PAGEBREAK(doc)

    # ==== Chapter 3 ====================================================
    H2(doc, "CHAPTER 3: Conclusion and Results")
    P(doc,
        "This chapter summarises the measurable outcomes of the "
        "author's contribution to DentyHub at the conclusion of GP2 "
        "and identifies the personal and technical results most "
        "worth carrying forward.")
    P(doc,
        "At submission, DentyHub is a four-role web platform with "
        "both a local development build and a production-shape "
        "deployment configuration (Vercel for the frontend, Fly.io "
        "for the backend, and Neon for the managed PostgreSQL "
        "instance). The measurable outcomes most directly "
        "attributable to the author's contribution are summarised "
        "below.")
    B(doc, [
        "Thirty-two functional features are documented in Section "
        "2.1 of the group report, with sixty functional requirements "
        "and thirty-four non-functional requirements mapped feature-"
        "by-feature in Section 2.2.",
        "Eleven NestJS controllers expose one hundred and nineteen "
        "REST endpoints. All endpoints are guarded by the "
        "JwtAuthGuard at the controller layer; elevated endpoints "
        "additionally enforce the AdminGuard.",
        "The Prisma schema comprises forty-one models, eighteen "
        "enumerations, and twelve append-only migrations. Foreign-"
        "key referential integrity is enforced at the database "
        "level.",
        "The arcade subsystem comprises six games with per-level "
        "leaderboards, sticky sequential unlock semantics, server-"
        "side enforcement of the once-per-day attempt rule, and a "
        "separate daily trivia quiz module.",
        "The Healthy Smile Streak feature implements Asia/Amman "
        "date-key arithmetic, automatic badge awards at four "
        "milestones (3, 7, 30, and 100 consecutive days), and a "
        "three-metric leaderboard.",
        "The internationalisation dictionary contains approximately "
        "one thousand one hundred keys distributed across English "
        "and Arabic. All patient-facing screens render with "
        "complete right-to-left layout in Arabic, as verified by "
        "TC-28 in Section 5.2 of the group report.",
        "All four user roles support light and dark themes; all "
        "patient-facing screens render at a 360-pixel viewport "
        "without horizontal scrolling; the settings panel exposes a "
        "reduced-motion preference compliant with WCAG 2.1 SC "
        "2.3.3.",
        "All thirty scenario tests pass against the freeze build. "
        "Login round-trip latency at the 95th percentile is below "
        "the 250-millisecond target specified in NFR-2; dashboard "
        "first contentful paint is below the 1500-millisecond "
        "target specified in NFR-7 against a simulated 4G "
        "connection.",
    ])
    P(doc,
        "The personal result the author considers most significant "
        "is not any individual feature but the overall maintainability "
        "of the codebase. Module docstrings are present throughout "
        "the backend; schema decisions are explained in inline "
        "comments; the migration sequence forms a coherent narrative "
        "of the system's evolution; the internationalisation "
        "dictionary uses a clear key namespace convention; and the "
        "engine contracts are explicit. The intended outcome is that "
        "the gap between the present implementation and a "
        "production-deployable system is substantially smaller than "
        "would have been the case had short-term deadline pressure "
        "been permitted to compromise these qualities.")
    P(doc,
        "Looking beyond the present project, the principles the "
        "author intends to carry forward are: prioritising data-"
        "model design over user-interface design; encoding process "
        "rules in written form rather than relying on team memory; "
        "exercising explicit discipline regarding scope; and "
        "establishing modular boundaries that may not provide "
        "immediate benefit but reduce future maintenance cost. The "
        "author acknowledges that these principles were not applied "
        "with perfect consistency during the present project; the "
        "intent for future work is to apply them more "
        "systematically.")
    P(doc,
        "The author gratefully acknowledges Dr. Zakarea M. Al "
        "Shara'a for his sustained supervision throughout both "
        "academic cycles; the four team members - Suhib Naser "
        "Rizek Hawwari, Nabeel Fadl Nabeel Aldalqamoni, Batool Amin "
        "Ali Allatayfeh, and Ragd Sami Mousa Al Qawasmi - for "
        "their respective contributions, which are documented in "
        "their individual reflection reports; the friends from the "
        "Faculty of Dentistry who provided domain expertise that "
        "shaped the data model; and the open-source maintainers "
        "of the NestJS, Prisma, Next.js, React, Tailwind, "
        "three.js, and react-three-fiber projects upon which this "
        "work depends.")
    return doc


# ---------------------------------------------------------------------------
# Teammate reports — academic style
# ---------------------------------------------------------------------------

TEAMMATE_SPECS = {
    "suhib": {
        "name": "Suhib Naser Rizek Hawwari",
        "student_id": "154860",
        "task_title": (
            "Administrative create-read-update-delete endpoints for "
            "users, rotation plans, and doctor groups"
        ),
        "contribution_paragraph": (
            "The author's principal contribution to DentyHub comprised "
            "the implementation of the backend create-read-update-"
            "delete (CRUD) endpoints supporting the administrative "
            "data-management surfaces. Specifically, the author "
            "authored the endpoints for the four role-typed user "
            "directories (patient, doctor, supervisor, and "
            "administrator), for rotation plans and their per-day "
            "assignments, and for doctor groups and group memberships. "
            "These endpoints constitute the data-management layer that "
            "transforms the administrator planning workspace from a "
            "read-only directory into an operational management tool."
        ),
        "what_worked": (
            "The decision to follow the existing controller-per-domain "
            "convention rather than to introduce a parallel structure "
            "produced consistent and predictable interfaces. Each new "
            "endpoint resides within the appropriate domain controller "
            "alongside existing peer endpoints, applies the JwtAuthGuard "
            "and the appropriate role guard at the controller layer, "
            "and validates incoming payloads through class-validator "
            "data transfer objects with whitelist enforcement enabled. "
            "This consistency simplified consumption by the frontend "
            "and reduced the cognitive overhead required during code "
            "review."
        ),
        "what_to_improve": (
            "In retrospect, the data transfer objects should have been "
            "specified before the controller signatures. The chosen "
            "order - controllers first, with data transfer objects "
            "back-filled - resulted in two validation gaps that were "
            "subsequently identified during the manual scenario-test "
            "pass. Adopting a data-transfer-object-first design would "
            "have prevented these gaps and would have permitted "
            "TypeScript type-checking to enforce consistency between "
            "the payload structure and the controller interface."
        ),
        "other_lesson": (
            "The implementation of CRUD endpoints surfaces a "
            "significant design decision at each entity: whether to "
            "support hard deletion, soft deletion, or deactivation. "
            "The decision depends on the entity's referential context. "
            "A clinic must be soft-deleted because referencing "
            "appointment records would otherwise be orphaned; a "
            "rotation-plan day may be hard-deleted because the parent "
            "plan record remains intact; a doctor group requires "
            "deactivation rather than deletion because audit "
            "considerations require the historical record to be "
            "preserved. The lesson is that CRUD design is not "
            "mechanical but requires per-entity consideration of "
            "referential semantics."
        ),
        "specialised": (
            "Backend data-management endpoints and data transfer "
            "object specification"
        ),
        "specialised_body": (
            "The author implemented the backend endpoints consumed by "
            "the administrator planning surface. These include the "
            "user directory CRUD across the four user roles "
            "(comprising creation, search and filter, block and "
            "unblock, re-approval of previously rejected accounts, "
            "and the freeze and unfreeze operations on doctor "
            "accounts); the rotation plan CRUD (comprising creation, "
            "editing, and deletion of the plan record and its "
            "associated per-day assignments); and the doctor group "
            "CRUD (comprising group creation, member addition and "
            "removal, supervisor addition and removal, and the "
            "processing of group join requests). Each endpoint "
            "validates incoming requests via a class-validator data "
            "transfer object configured with whitelist and "
            "forbidNonWhitelisted enforcement enabled; resides behind "
            "the appropriate authorisation guard; and executes within "
            "a Prisma transaction when the operation requires "
            "modification of multiple records. The bulk semester-"
            "cohort progression endpoint is illustrative of the latter "
            "pattern: a single administrative action rewrites the "
            "semester pointer for every doctor within a group "
            "atomically."
        ),
        "conclusion_paragraph": (
            "The endpoints authored by the author constitute the "
            "data-management layer upon which the administrator "
            "workspace depends. They appear in the Appendix C API "
            "catalogue of the group report under the AdminController "
            "and SupervisorController owners, and are validated by "
            "test cases TC-05 (administrator approval of a doctor "
            "account) and TC-30 (administrator use of the Add-case "
            "deep-link shortcut). While these endpoints are not "
            "individually visible to end users, they constitute the "
            "infrastructure that renders the administrator role "
            "operationally useful."
        ),
    },
    "nabeel": {
        "name": "Nabeel Fadl Nabeel Aldalqamoni",
        "student_id": "158189",
        "task_title": (
            "Use case analysis, system architecture, leaderboard "
            "implementation, and documentation figure evidence"
        ),
        "contribution_paragraph": (
            "The author's principal contributions to DentyHub "
            "comprised four interrelated work products: the use case "
            "analysis defining the capabilities available to each user "
            "role; the high-level system architecture diagram "
            "documenting the four-tier structure of the platform; the "
            "implementation of the per-game per-level arcade "
            "leaderboard and the Healthy Smile Streak leaderboard; and "
            "the capture and integration of the thirty-eight "
            "screenshots evidencing Chapter 4 of the group report."
        ),
        "what_worked": (
            "The decision to perform use case analysis prior to "
            "architectural design proved productive. The use case "
            "diagram (Figure 3.2 of the group report) enumerates the "
            "capabilities available to each of the four user roles, "
            "and this enumeration informed both the functional-"
            "features table in Section 2.1 and the layered "
            "architecture diagram in Figure 3.1. Investing analytical "
            "effort in the use cases before architectural design "
            "ensured that the architecture decomposition reflected "
            "actual system capabilities rather than speculative "
            "ones."
        ),
        "what_to_improve": (
            "The leaderboard user interface was constructed prior to "
            "the finalisation of the ranking semantics, which "
            "necessitated two subsequent layout revisions when the "
            "metric definitions were clarified. In retrospect, the "
            "metric definitions - the distinction between current "
            "streak, best-ever streak, and cumulative points; and "
            "between per-game per-level ranking and combined ranking - "
            "should have been documented before the user interface "
            "design commenced."
        ),
        "other_lesson": (
            "Documentation figures bear a significant share of the "
            "communicative burden of the group report. A reviewer who "
            "skims the prose nevertheless examines the diagrams and "
            "screenshots, and inconsistent or low-quality figures "
            "undermine the credibility of the surrounding claims even "
            "when those claims are accurate. Consistent figure "
            "treatment - uniform viewport, consistent thematic "
            "grouping, consistent border and caption styling - "
            "produces a disproportionate return on investment."
        ),
        "specialised": (
            "Systems analysis, leaderboard implementation, and "
            "documentation evidence"
        ),
        "specialised_body": (
            "The author authored the use case diagram (Figure 3.2 of "
            "the group report), which presents all four user roles "
            "(patient, student doctor, clinical supervisor, and "
            "administrator) within a single system boundary, with "
            "role-specific use case packages and a cross-cutting "
            "package for use cases performed by all authenticated "
            "roles (in-application notifications, peer-to-peer chat, "
            "language switching, and profile editing). The author "
            "additionally produced the high-level system architecture "
            "diagram (Figure 3.1 of the group report), which "
            "presents the four-tier structure of the platform - the "
            "client tier, the presentation tier (Next.js 16), the "
            "application tier (NestJS 11), and the data tier "
            "(Prisma 5.22 over PostgreSQL 16) - together with the "
            "operations subsystem (environment configuration and "
            "background job execution)."
        ),
        "specialised_body_extra": (
            "The author implemented the leaderboard surfaces visible "
            "to both patients and doctors. The arcade leaderboard "
            "comprises a tab control selecting among the six arcade "
            "games, a level filter dropdown selecting among Levels 1 "
            "through 11, rank rendering with podium emoji "
            "decoration for the top three positions and numeric "
            "ranks below, and a highlight treatment indicating the "
            "calling user's row. The streak leaderboard, accessible "
            "via the streak tab of the patient leaderboard "
            "switcher, comprises a three-metric switcher (current "
            "streak, best-ever streak, and cumulative points) and "
            "a hero card displaying the calling user's rank."
        ),
        "specialised_body_third": (
            "The author captured the thirty-eight screenshots "
            "evidencing Chapter 4 of the group report. The "
            "screenshots are organised into thematic groups: public "
            "and authentication screens (Figures 4.1 through 4.3); "
            "patient surfaces (Figures 4.4 through 4.9); doctor "
            "surfaces (Figures 4.10 through 4.14); supervisor "
            "surfaces (Figures 4.15 through 4.17); administrator "
            "surfaces (Figures 4.18 through 4.26); arcade surfaces "
            "(Figures 4.27 through 4.34); and Arabic and "
            "accessibility surfaces (Figures 4.35 through 4.38). "
            "Each screenshot was captured at a consistent viewport, "
            "named to correspond to its caption, and inserted into "
            "the appropriate placeholder cell within the document "
            "without disrupting the surrounding prose."
        ),
        "conclusion_paragraph": (
            "The two most visually prominent figures in the group "
            "report (Figure 3.1, the system architecture, and "
            "Figure 3.2, the use case analysis) were authored by "
            "the author, as were the leaderboard surfaces visible "
            "to all users and the thirty-eight screenshots "
            "evidencing Chapter 4. The use case analysis "
            "additionally informed the ordering of the functional-"
            "features table in Section 2.1, ensuring that the "
            "analytical work contributed to the requirements "
            "engineering chapter as well as to the design chapter."
        ),
    },
    "batool": {
        "name": "Batool Amin Ali Allatayfeh",
        "student_id": "152284",
        "task_title": (
            "Authentication infrastructure, user profile management, "
            "and identity-related surfaces"
        ),
        "contribution_paragraph": (
            "The author's principal contribution to DentyHub "
            "comprised the implementation of the authentication "
            "infrastructure - registration, login, session "
            "establishment, and password management - together with "
            "the user profile editing surfaces and the user profile "
            "reporting workflow. The authentication infrastructure "
            "encompasses password hashing via the bcryptjs library, "
            "JSON Web Token issuance and verification via the "
            "passport-jwt library, the placement of authorisation "
            "guards (JwtAuthGuard and AdminGuard) at the controller "
            "layer, and the migration from the legacy header-based "
            "actor identification pattern to the present token-"
            "based pattern."
        ),
        "what_worked": (
            "The decision to place the AdminGuard at the controller "
            "layer rather than at the route table proved sound. "
            "Route-table guarding has the advantage of declarative "
            "specification but presents a failure mode in which a "
            "newly added controller silently leaks data when the "
            "corresponding route-table entry is omitted. Controller-"
            "layer guarding co-locates the authorisation check with "
            "the data it protects, eliminating the failure mode at "
            "the cost of slightly more verbose specification. NFR-6 "
            "of the group report captures this design decision."
        ),
        "what_to_improve": (
            "The failed-login lockout behaviour (NFR-13: five failed "
            "attempts within five minutes locks the account for five "
            "minutes) should have been verified by an automated test "
            "prior to manual scenario testing. The chosen sequence "
            "of implementation followed by manual verification "
            "permitted a clock-skew edge case to remain undetected "
            "until late in GP2. Test-first development would have "
            "surfaced the defect immediately."
        ),
        "other_lesson": (
            "Authentication implementation is principally concerned "
            "with precise rejection of unauthorised access. Each "
            "additional endpoint introduces an additional opportunity "
            "for unauthorised access if the appropriate guard is "
            "omitted; each exception to the standard authorisation "
            "policy (impersonation for support purposes, elevated "
            "administrator actions, and the partner-doctor case in "
            "which one doctor may access another doctor's "
            "appointment data) represents an additional row in an "
            "implicit permissions matrix. The lesson is that "
            "maintaining this permissions matrix as explicit "
            "documentation is preferable to maintaining it as "
            "implicit team knowledge."
        ),
        "specialised": (
            "Authentication, identity management, and user profile "
            "infrastructure"
        ),
        "specialised_body": (
            "The author implemented the authentication module "
            "located at backend/src/auth. The AuthService exposes "
            "the registration, login, change-password, update-"
            "profile, and resend-credential-request operations. The "
            "AuthController exposes the corresponding HTTP "
            "endpoints. The JwtStrategy and JwtAuthGuard implement "
            "the passport-jwt integration. The AdminGuard "
            "implements the elevated authorisation check. The "
            "@CurrentUser decorator and the JwtPayload type "
            "support the controller-layer access pattern. "
            "Passwords are hashed via bcryptjs at a cost factor of "
            "10; the JSON Web Token is signed with the JWT_SECRET "
            "environment variable and expires after seven days. "
            "The failed-login lockout (five failed attempts within "
            "five minutes locking the account for five minutes) "
            "is enforced at the server side. The migration from "
            "the legacy x-actor-* header pattern to the present "
            "token-based pattern was completed in a single commit "
            "with end-to-end replacement of the authentication "
            "workflow and complete commit-message documentation."
        ),
        "specialised_body_extra": (
            "The user profile surfaces built upon the authentication "
            "identity comprise the profile editing workflow (name, "
            "telephone number, biographical text, and profile "
            "picture; with administrator-only fields locked); the "
            "profile popup (the public profile view displayed when "
            "any user clicks another user's avatar); and the "
            "UserProfileReport queue, which permits any user to "
            "file a misconduct report against another user for "
            "administrator review. The profile data structures are "
            "centralised within frontend/src/features/profiles so "
            "that all four role-specific profile views consume the "
            "same underlying component rather than re-implementing "
            "the data presentation independently."
        ),
        "conclusion_paragraph": (
            "The authentication infrastructure implemented by the "
            "author constitutes the gateway through which all "
            "DentyHub requests pass. It is documented in the "
            "group report as feature F1 (account registration and "
            "JWT login); is subject to non-functional requirements "
            "NFR-1 through NFR-5 of the security cluster in "
            "Section 2.2; is validated by test cases TC-01 "
            "through TC-05 in Section 5.2; and is enumerated row "
            "by row in Appendix C for every /auth/* endpoint. The "
            "user profile surfaces (features F4, F28, and F30) "
            "extend this same identity infrastructure. No "
            "subsequent component of the system functions "
            "correctly if this layer is incorrect, a property "
            "that justified the deliberate pace of "
            "implementation."
        ),
    },
    "ragd": {
        "name": "Ragd Sami Mousa Al Qawasmi",
        "student_id": "162083",
        "task_title": (
            "Clinical cases domain - semester clinic cases, per-doctor "
            "case progress, and integration with the booking and "
            "supervisor decision workflows"
        ),
        "contribution_paragraph": (
            "The author's principal contribution to DentyHub "
            "comprised the design and implementation of the clinical "
            "cases domain. This domain encompasses the "
            "SemesterClinicCase model and its associated catalogue "
            "management; the DoctorClinicCaseProgress denormalised "
            "rollup model tracking each doctor's progress toward the "
            "graduation requirement for each case category; the "
            "integration of the cases domain with the doctor groups "
            "and semester rotation infrastructure; and the case-"
            "state transitions triggered by the supervisor decision "
            "workflow."
        ),
        "what_worked": (
            "The decision to consult a domain expert from the "
            "Faculty of Dentistry before drafting the schema "
            "produced significant design value. The clarification "
            "that a redo decision implies treatment of a different "
            "patient rather than re-treatment of the same patient "
            "shaped the entire data model: it necessitated the "
            "DoctorClinicCaseProgress denormalised table rather than "
            "a derived view, because a redo on one appointment must "
            "not credit a separate appointment touching the same "
            "case. Domain consultation prior to implementation "
            "produced a more accurate model than would have resulted "
            "from inference alone."
        ),
        "what_to_improve": (
            "The initial implementation attempted to derive per-"
            "doctor case progress dynamically from the appointment "
            "and report records, on the assumption that "
            "denormalisation should be avoided where possible. While "
            "this approach produced correct results, the five-table "
            "join required to compute per-case progress proved slow "
            "enough that the doctor dashboard's case-progress "
            "component exhibited visible loading delay. The "
            "subsequently adopted denormalised DoctorClinicCaseProgress "
            "table with transactional updates triggered by the "
            "supervisor decision constitutes the appropriate "
            "solution. In retrospect, the derived-view "
            "implementation should have been bypassed in favour of "
            "the denormalised approach from the outset."
        ),
        "other_lesson": (
            "Domain modelling is more accurately characterised as "
            "constraint specification than as field specification. "
            "The UNIQUE constraint on (doctorId, clinicCaseId) "
            "applied to DoctorClinicCaseProgress is the property "
            "that permits the question \"has this doctor completed "
            "this case category?\" to be answered by a single "
            "Boolean lookup rather than by a count operation "
            "combined with a comparison. The UNIQUE constraint on "
            "(semesterId, clinicId, title) applied to "
            "SemesterClinicCase prevents the duplicate listing of "
            "the same case when an administrator re-creates it. "
            "Specifying these constraints at the schema level "
            "obviated substantial defensive code at the application "
            "level."
        ),
        "specialised": (
            "Clinical cases domain modelling and per-doctor "
            "graduation progress tracking"
        ),
        "specialised_body": (
            "The author designed and implemented the clinical "
            "cases domain comprising three principal models. The "
            "SemesterClinicCase model represents the bookable case "
            "category, scoped to a particular semester at a "
            "particular clinic, with an integer requiredCount "
            "field encoding the graduation requirement. The "
            "DoctorClinicCaseProgress model represents the per-"
            "(doctor, case) rollup, comprising a status "
            "enumeration (OPEN, ASSISTED, COMPLETED) and back-"
            "pointer fields lastReportId and lastAppointmentId. "
            "The Appointment model was extended to reference a "
            "SemesterClinicCase, ensuring that a patient cannot "
            "reserve a slot without simultaneously specifying a "
            "case category."
        ),
        "specialised_body_extra": (
            "The author additionally implemented the case-state "
            "transitions triggered by the three supervisor decision "
            "outcomes. The APPROVE decision updates "
            "DoctorClinicCaseProgress to the COMPLETED status "
            "within the same Prisma transaction as the CaseReport "
            "status update, ensuring that the graduation-credit "
            "record cannot diverge from the supervisor's authoritative "
            "record. The REDO decision leaves the progress record "
            "at the OPEN status and invalidates the underlying "
            "appointment, requiring the doctor to repeat the case "
            "category with a different patient on a fresh "
            "appointment. The NEEDS_EDIT decision affects neither "
            "the progress record nor the appointment binding; the "
            "report is unlocked for resubmission with the same "
            "patient and doctor. These transitions are documented "
            "in Figure 3.6 (the appointment state diagram) and "
            "Figure 3.5 (the supervisor decision sequence diagram) "
            "of the group report."
        ),
        "specialised_body_third": (
            "The author additionally implemented the integration of "
            "the cases domain with the administrator surfaces, "
            "comprising the administrator Cases page (which "
            "presents the semester catalogue with per-clinic "
            "grouping, soft-delete and restore actions, and the "
            "per-doctor progress override panel) and the bulk "
            "semester-cohort progression action (which advances "
            "every doctor in a DoctorGroup to the subsequent year "
            "within a single Prisma transaction). The bulk "
            "progression action exemplifies the all-or-nothing "
            "semantics required by the operation: either every "
            "doctor in the cohort advances or no doctor advances; "
            "no intermediate state is permitted."
        ),
        "conclusion_paragraph": (
            "The clinical cases domain implemented by the author "
            "constitutes the integrative element between the "
            "academic and clinical aspects of the system. It is "
            "documented in the group report as features F5 "
            "(appointment request with clinic case selection), "
            "F11 (administrator planning workspace), F26 (per-"
            "doctor case progress override), and F27 (bulk "
            "semester progression); is governed by functional "
            "requirements FR-13, FR-21, FR-25 through FR-27, FR-"
            "49, and FR-50 of Section 2.2; and is validated by "
            "test cases TC-12 (APPROVE updates progress) and TC-"
            "14 (REDO does not credit). The atomic-decision "
            "guarantee provided by this domain represents the "
            "architectural property that distinguishes DentyHub "
            "from a generic appointment management system."
        ),
    },
}


def teammate_report(spec):
    doc = new_doc()
    TITLE_PAGE(doc, spec["name"], spec["student_id"])
    UNDERTAKING(doc, spec["name"])
    CONTENTS(doc)

    # ==== Chapter 1 ====================================================
    H2(doc, "CHAPTER 1: Lessons Learned")
    P(doc,
        "This chapter documents the lessons derived from the author's "
        "participation in the DentyHub project. The reflection is "
        "organised into three categories: strategies and processes "
        "that contributed to the success of the author's "
        "deliverables; areas in which the author's approach could be "
        "improved; and broader observations gained through "
        "participation in a multi-person software engineering "
        "project.")

    H3(doc, "1.1 Strategies and Processes that led to Success")
    P(doc, spec["contribution_paragraph"])
    P(doc, spec["what_worked"])
    P(doc,
        "Beyond the contribution described above, the author "
        "attended weekly supervisory check-ins with Dr. Zakarea M. "
        "Al Shara'a and participated in code review on pull "
        "requests covering work surfaces adjacent to the author's "
        "own. Observing the project lead's discipline regarding "
        "data-model design and migration management contributed to "
        "the author's understanding of professional software "
        "engineering practice, which the author characterises as "
        "the consistent application of small habits rather than the "
        "occasional execution of large heroic actions.")

    H3(doc, "1.2 Areas of potential improvement")
    P(doc, spec["what_to_improve"])
    P(doc,
        "Additionally, the author observes that engagement with the "
        "wider architectural decisions of the project could have "
        "been more substantial. The focused scope of the author's "
        "contribution during GP2 was appropriate to the time "
        "available and to the depth required for the chosen "
        "specialisation; however, broader participation in "
        "architectural discussion would have enhanced the author's "
        "confidence when addressing system-level questions. In "
        "future projects, the author intends to combine focused "
        "specialisation with one stretch contribution at the "
        "architectural level.")

    H3(doc, "1.3 Other lessons learned")
    P(doc, spec["other_lesson"])
    P(doc,
        "The author additionally observes that working on a "
        "project addressing a tangible social purpose - the "
        "provision of supervised dental care to patients who would "
        "otherwise lack access to it - differs qualitatively from "
        "the abstraction of standard coursework. The user is no "
        "longer an abstract specification but a person whose "
        "interaction with the system has practical consequences. "
        "Conversations with practitioners from the Faculty of "
        "Dentistry clarifying the operational meaning of clinical "
        "concepts (for example, the precise definition of a redo "
        "request) made the data-model decisions more concrete than "
        "they would otherwise have been.")
    PAGEBREAK(doc)

    # ==== Chapter 2 ====================================================
    H2(doc, "CHAPTER 2: Individual Contributions")
    P(doc,
        "This chapter characterises the author's contribution to "
        "DentyHub. Qualitative claims are cross-referenced to the "
        "corresponding sections, figures, features, requirements, "
        "and test cases of the group report.")

    H3(doc, "2.1 Amount of Work")
    P(doc,
        f"The author's contribution to DentyHub comprised a "
        f"deliverable scoped to be defensible at the level of "
        f"detail required by the project viva. The deliverable - "
        f"\"{spec['task_title']}\" - was executed against the "
        f"project freeze build during the final two weeks of GP2 "
        f"and is documented in the artefacts contributed to the "
        f"project repository.")

    H3(doc, "2.2 Quality of Work")
    P(doc,
        "The quality standard applied to this deliverable required "
        "that each observation produced by the author be "
        "reproducible by a different team member from the same "
        "starting state, and that each defect or design decision "
        "filed be sufficiently precise that the receiving team "
        "member could act upon it without further consultation. "
        "Concretely, the author recorded the actual behaviour "
        "against the expected behaviour for each case considered, "
        "and re-validated the failing cases after the corresponding "
        "fixes were applied. The discipline of re-validation, "
        "while readily omitted under deadline pressure, prevents "
        "regression and was treated as non-optional in this "
        "deliverable.")

    H3(doc, "2.3 Specialized Contributions")
    P(doc, f"Specialisation: {spec['specialised']}.", bold=True)
    P(doc, spec["specialised_body"])
    if spec.get("specialised_body_extra"):
        P(doc, spec["specialised_body_extra"])
    if spec.get("specialised_body_third"):
        P(doc, spec["specialised_body_third"])

    H3(doc, "2.4 Initiative and Reliability")
    P(doc,
        "The author proposed the scope of the present deliverable, "
        "executed it without requiring detailed guidance from the "
        "project lead, and reported completion through the "
        "project's standard communication channel on the same day "
        "as the freeze pass concluded. The deliverable was "
        "completed within the originally committed timeframe.")
    P(doc,
        "Reliability is documented by the sign-off message posted "
        "at the conclusion of the freeze pass and by the absence "
        "of subsequent rework requirements arising from the "
        "deliverable.")
    PAGEBREAK(doc)

    # ==== Chapter 3 ====================================================
    H2(doc, "CHAPTER 3: Conclusion and Results")
    P(doc, spec["conclusion_paragraph"])
    P(doc,
        "Considered in the wider context of the project, DentyHub "
        "constitutes a four-role platform encompassing the full "
        "case lifecycle from patient request to supervisor-"
        "approved case report, augmented by a patient engagement "
        "layer (streak, badges, six arcade games, and a daily "
        "trivia quiz) and complete English and Arabic "
        "localisation with right-to-left presentation on every "
        "patient-facing screen. The author's deliverable "
        "integrates into this wider system as one verifiable "
        "thread of contribution; the full system documentation "
        "appears in Chapters 1 through 5 and the twelve "
        "appendices of the group report.")
    P(doc,
        "The most significant result the author identifies for "
        "personal carry-forward is the experience of taking sole "
        "responsibility for a defined component of a real system. "
        "This experience differs from coursework in that the "
        "deliverable will be examined by a panel and may "
        "subsequently be used in production. Awareness of this "
        "downstream consumption altered the author's working "
        "standard during execution and constitutes the principal "
        "professional skill developed through participation in "
        "the project.")
    P(doc,
        "The author gratefully acknowledges Dr. Zakarea M. Al "
        "Shara'a for the project supervision; Omar Mohammed "
        "Subhi Al Shamali for the technical leadership of the "
        "team and the codebase; and the remaining team members "
        "for their respective contributions, which are "
        "documented in their individual reflection reports.")
    return doc


# ---------------------------------------------------------------------------
# Save with fallback filenames so Word lock does not block regeneration
# ---------------------------------------------------------------------------


def save_doc(doc, fname):
    candidates = [fname,
                  fname.replace(".docx", "_v2.docx"),
                  fname.replace(".docx", "_v3.docx")]
    for path in candidates:
        try:
            doc.save(path)
            print(f"Wrote {path}")
            return
        except PermissionError:
            continue
    print(f"All candidates locked for {fname}.")


save_doc(omar_report(), "DentyHub_Individual_Omar_Al_Shamali.docx")
save_doc(teammate_report(TEAMMATE_SPECS["suhib"]), "DentyHub_Individual_Suhib_Hawwari.docx")
save_doc(teammate_report(TEAMMATE_SPECS["nabeel"]), "DentyHub_Individual_Nabeel_Aldalqamoni.docx")
save_doc(teammate_report(TEAMMATE_SPECS["batool"]), "DentyHub_Individual_Batool_Allatayfeh.docx")
save_doc(teammate_report(TEAMMATE_SPECS["ragd"]), "DentyHub_Individual_Ragd_Al_Qawasmi.docx")
